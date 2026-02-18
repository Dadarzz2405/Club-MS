"""
Rohis Management System — REST API Backend
All routes return JSON. Designed for a Next.js (or any SPA) frontend.

Auth strategy: session-based (Flask-Login + cookie).
For cross-origin deployments set FRONTEND_ORIGIN in .env and ensure
the frontend sends credentials (credentials: 'include' in fetch).
"""

from utils import can_mark_attendance, is_core_user
from flask import Flask, request, jsonify, abort, Response
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, login_user, logout_user,
    login_required, current_user,
)
from flask_bcrypt import Bcrypt
from flask_cors import CORS
from flask_migrate import Migrate
import os
from dotenv import load_dotenv
from models import (
    Pic, db, User, Session, Attendance, Notulensi,
    Division, JadwalPiket, PiketAssignment, EmailReminderLog, SessionPIC,
)
from datetime import datetime, date, timezone, timedelta, time
from ummalqura.hijri_date import HijriDate
import json
from werkzeug.utils import secure_filename
from ai import call_chatbot_groq
import csv
from io import TextIOWrapper, StringIO, BytesIO
from docx import Document
from summarizer import summarize_notulensi
from sqlalchemy.exc import IntegrityError
from email_service import get_email_service
import logging
import os
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {"png", "jpg", "jpeg", "webp"}

# ---------------------------------------------------------------------------
# App & extension init
# ---------------------------------------------------------------------------
app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY')
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL')
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
db.init_app(app)
bcrypt = Bcrypt(app)
migrate = Migrate(app, db)

login_manager = LoginManager()
login_manager.init_app(app)

# Return 401 JSON instead of redirecting to a login view
@login_manager.unauthorized_handler
def unauthorized():
    return jsonify({"success": False, "error": "unauthorized", "message": "Login required"}), 401


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def serialize_user(user, include_email=True):
    """Return a safe dict representation of a User."""
    data = {
        "id": user.id,
        "name": user.name,
        "role": user.role,
        "class_name": user.class_name,
        "can_mark_attendance": user.can_mark_attendance,
        "must_change_password": user.must_change_password,
        "pic_id": user.pic_id,
        "pic_name": user.pic.name if user.pic else None,
        "profile_picture_url": f"/api/profile/picture/{user.id}",
    }
    if include_email:
        data["email"] = user.email
    return data


def serialize_session(s):
    return {
        "id": s.id,
        "name": s.name,
        "date": s.date,
        "is_locked": s.is_locked,
        "session_type": s.session_type,
        "description": s.description,
        "created_at": s.created_at.isoformat() if s.created_at else None,
        "assigned_pics": [{"id": p.id, "name": p.name} for p in s.assigned_pics],
        "attendance_count": len(s.attendances),
    }


def serialize_pic(pic):
    return {
        "id": pic.id,
        "name": pic.name,
        "description": pic.description,
        "created_at": pic.created_at.isoformat() if pic.created_at else None,
        "member_count": len(pic.members),
        "members": [{"id": m.id, "name": m.name} for m in pic.members],
    }


def serialize_attendance(att):
    wib = timezone(timedelta(hours=7))
    return {
        "id": att.id,
        "session_id": att.session_id,
        "session_name": att.session.name if att.session else None,
        "session_date": att.session.date if att.session else None,
        "user_id": att.user_id,
        "status": att.status,
        "attendance_type": att.attendance_type,
        "timestamp": att.timestamp.astimezone(wib).isoformat() if att.timestamp else None,
    }


def serialize_notulensi(note):
    return {
        "id": note.id,
        "session_id": note.session_id,
        "session_name": note.session.name if note.session else None,
        "session_date": note.session.date if note.session else None,
        "content": note.content,
        "created_at": note.created_at.isoformat() if note.created_at else None,
        "updated_at": note.updated_at.isoformat() if note.updated_at else None,
    }


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------

@app.route("/health")
def health_check():
    return jsonify({
        'status': 'ok',
        'service': 'Rohis Attendance System',
        'timestamp': datetime.utcnow().isoformat(),
        'message': 'App is awake and running'
    })
#Login, no need comment actually
@app.route('/login', methods=['GET', 'POST'])
def login():
    data = request.get_json() or {}
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")

    if not email or not password:
        return jsonify({"success": False, "message": "Email and password are required"}), 400

    user = User.query.filter_by(email=email).first()
    if not user or not bcrypt.check_password_hash(user.password, password):
        return jsonify({"success": False, "message": "Invalid email or password"}), 401

    login_user(user)
    return jsonify({
        "success": True,
        "user": serialize_user(user),
        "must_change_password": user.must_change_password,
    })


@app.route("/api/auth/logout", methods=["POST"])
@login_required
def logout():
    logout_user()
    return jsonify({"success": True, "message": "Logged out"})


@app.route("/api/auth/me")
@login_required
def me():
    return jsonify({"success": True, "user": serialize_user(current_user)})


# ===========================================================================
# PROFILE
# ===========================================================================

@app.route("/api/profile", methods=["PUT"])
@login_required
def update_profile():
    data = request.get_json() or {}
    username = data.get("username", "").strip()
    password = data.get("password", "")

    if username:
        existing = User.query.filter_by(username=username).first()
        if existing and existing.id != current_user.id:
            return jsonify({"success": False, "message": "Username already taken"}), 409
        current_user.username = username

    if password:
        current_user.password = bcrypt.generate_password_hash(password).decode("utf-8")

    db.session.commit()
    return jsonify({"success": True, "message": "Profile updated", "user": serialize_user(current_user)})


@app.route("/api/profile/password", methods=["PUT"])
@login_required
def change_password():
    data = request.get_json() or {}
    old_password = data.get("old_password", "")
    new_password = data.get("new_password", "")
    confirm_password = data.get("confirm_password", "")

    if not bcrypt.check_password_hash(current_user.password, old_password):
        return jsonify({"success": False, "message": "Incorrect current password"}), 400
    if new_password != confirm_password:
        return jsonify({"success": False, "message": "New passwords do not match"}), 400
    if len(new_password) < 6:
        return jsonify({"success": False, "message": "Password must be at least 6 characters"}), 400

    current_user.password = bcrypt.generate_password_hash(new_password).decode("utf-8")
    current_user.must_change_password = False
    db.session.commit()
    return jsonify({"success": True, "message": "Password updated successfully"})


@app.route("/api/profile/picture", methods=["POST"])
@login_required
def upload_pfp():
    file = request.files.get("pfp")
    if not file or not file.filename:
        return jsonify({"success": False, "message": "No file provided"}), 400

    if not allowed_file(file.filename):
        return jsonify({"success": False, "message": "Invalid file type. Allowed: png, jpg, jpeg, webp"}), 400

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 5 * 1024 * 1024:
        return jsonify({"success": False, "message": "File too large (max 5 MB)"}), 400

    current_user.profile_picture_data = file.read()
    current_user.profile_picture_filename = secure_filename(file.filename)
    db.session.commit()
    return jsonify({
        "success": True,
        "message": "Profile picture updated",
        "url": f"/api/profile/picture/{current_user.id}",
    })


@app.route("/api/profile/picture/<int:user_id>")
def serve_profile_picture(user_id):
    user = User.query.get_or_404(user_id)
    if user.profile_picture_data:
        filename = user.profile_picture_filename or "image.png"
        ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "png"
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg", "webp": "image/webp"}.get(ext, "image/png")
        return Response(user.profile_picture_data, mimetype=mime)

    default = os.path.join("static", "uploads", "profiles", "default.png")
    if os.path.exists(default):
        with open(default, "rb") as f:
            return Response(f.read(), mimetype="image/png")
    return jsonify({"error": "not_found"}), 404


# ===========================================================================
# MEMBERS
# ===========================================================================

@app.route("/api/members")
@login_required
def list_members():
    users = User.query.order_by(User.name).all()
    return jsonify({"success": True, "members": [serialize_user(u) for u in users]})


@app.route("/api/members", methods=["POST"])
@login_required
def add_member():
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    data = request.get_json() or {}
    name = data.get("name", "").strip()
    email = data.get("email", "").strip().lower()
    class_name = data.get("class_name") or None
    role = data.get("role", "member")

    if not name or not email:
        return jsonify({"success": False, "message": "Name and email are required"}), 400

    hashed = bcrypt.generate_password_hash("rohisnew").decode("utf-8")
    user = User(name=name, email=email, class_name=class_name, role=role, password=hashed)
    try:
        db.session.add(user)
        db.session.commit()
        return jsonify({"success": True, "message": f"Member {name} created", "member": serialize_user(user)}), 201
    except IntegrityError:
        db.session.rollback()
        return jsonify({"success": False, "message": "A user with that email already exists"}), 409


@app.route("/api/members/batch-add", methods=["POST"])
@login_required
def batch_add_members():
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    hashed = bcrypt.generate_password_hash("rohisnew").decode("utf-8")
    added, errors = 0, []

    def _add(name, email, class_name, role):
        nonlocal added
        if not name or not email:
            return
        email_l = email.strip().lower()
        if User.query.filter_by(email=email_l).first():
            errors.append(f"User with email {email_l} already exists")
            return
        try:
            db.session.add(User(
                name=name.strip(), email=email_l,
                class_name=class_name or None,
                role=role or "member", password=hashed,
            ))
            db.session.commit()
            added += 1
        except IntegrityError:
            db.session.rollback()
            errors.append(f"Failed to add {email_l}")

    # CSV file upload
    csv_file = request.files.get("csv_file")
    if csv_file and csv_file.filename:
        try:
            stream = TextIOWrapper(csv_file.stream, encoding="utf-8")
            for row in csv.reader(stream):
                if len(row) >= 2:
                    _add(row[0], row[1], row[2] if len(row) > 2 else None, row[3] if len(row) > 3 else "member")
        except Exception as e:
            errors.append(f"CSV parse error: {e}")

    # Bulk text (JSON body)
    bulk_text = (request.form.get("bulk_text") or "").strip()
    if not bulk_text and request.is_json:
        bulk_text = (request.get_json() or {}).get("bulk_text", "")
    for line in StringIO(bulk_text):
        parts = [p.strip() for p in line.strip().split(",")]
        if len(parts) >= 2:
            _add(parts[0], parts[1], parts[2] if len(parts) > 2 else None, parts[3] if len(parts) > 3 else "member")

    return jsonify({"success": True, "added": added, "errors": errors}), 201 if added else 200


@app.route("/api/members/batch-delete", methods=["POST"])
@login_required
def batch_delete_members():
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    data = request.get_json() or {}
    ids = data.get("ids", [])
    if not ids:
        return jsonify({"success": False, "message": "No member IDs provided"}), 400

    users_to_delete = User.query.filter(User.id.in_(ids)).all()
    if any(u.id == current_user.id for u in users_to_delete):
        return jsonify({"success": False, "message": "Cannot delete your own account"}), 400

    admin_count = User.query.filter_by(role="admin").count()
    removing_admins = sum(1 for u in users_to_delete if u.role == "admin")
    if admin_count - removing_admins < 1:
        return jsonify({"success": False, "message": "Cannot remove the last admin"}), 400

    deleted, failed = 0, []
    for u in users_to_delete:
        try:
            db.session.delete(u)
            db.session.commit()
            deleted += 1
        except Exception:
            db.session.rollback()
            failed.append(u.email)

    return jsonify({"success": True, "deleted": deleted, "failed": failed})


@app.route("/api/members/<int:user_id>", methods=["DELETE"])
@login_required
def delete_member(user_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    user = User.query.get_or_404(user_id)
    pic_id = request.form.get('pic_id')
    
    try:
        if pic_id and pic_id.strip():
            # Assign to PIC
            pic_id = int(pic_id)
            pic = Pic.query.get(pic_id)
            
            if not pic:
                flash(f'Invalid PIC selected', 'error')
                return redirect(url_for('member_management'))
            
            user.pic_id = pic_id
            flash(f'✅ {user.name} assigned to {pic.name}', 'success')
        else:
            # Remove PIC assignment
            if user.pic_id:
                old_pic = Pic.query.get(user.pic_id)
                user.pic_id = None
                flash(f'Removed {user.name} from {old_pic.name if old_pic else "PIC"}', 'info')
            else:
                flash(f'{user.name} has no PIC assignment', 'info')
        
        # Also update attendance permission if needed
        # user.can_mark_attendance = (user.pic_id is not None)
        
        db.session.commit()
        return jsonify({"success": True, "message": "Member deleted"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/members/<int:user_id>/role", methods=["PUT"])
@login_required
def change_member_role(user_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    user = User.query.get_or_404(user_id)
    data = request.get_json() or {}
    new_role = data.get("role")
    if not new_role:
        return jsonify({"success": False, "message": "Role is required"}), 400

    if user.role == "admin" and new_role != "admin":
        if User.query.filter_by(role="admin").count() <= 1:
            return jsonify({"success": False, "message": "Cannot remove the last admin's role"}), 400

    user.role = new_role
    db.session.commit()
    return jsonify({"success": True, "message": "Role updated", "member": serialize_user(user)})


@app.route("/api/members/<int:user_id>/pic", methods=["PUT"])
@login_required
def assign_member_pic(user_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    user = User.query.get_or_404(user_id)
    data = request.get_json() or {}
    pic_id = data.get("pic_id")  # None / null → remove assignment

    if pic_id:
        pic = Pic.query.get(pic_id)
        if not pic:
            return jsonify({"success": False, "message": "Invalid PIC"}), 404
        user.pic_id = pic_id
        message = f"{user.name} assigned to {pic.name}"
    else:
        user.pic_id = None
        message = f"PIC assignment removed from {user.name}"

    db.session.commit()
    return jsonify({"success": True, "message": message, "member": serialize_user(user)})


@app.route("/api/members/<int:user_id>/attendance-permission", methods=["PUT"])
@login_required
def toggle_attendance_permission(user_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    user = User.query.get_or_404(user_id)
    data = request.get_json() or {}
    # Accept explicit bool or toggle
    if "can_mark" in data:
        user.can_mark_attendance = bool(data["can_mark"])
    else:
        user.can_mark_attendance = not user.can_mark_attendance

    db.session.commit()
    return jsonify({
        "success": True,
        "can_mark_attendance": user.can_mark_attendance,
        "member": serialize_user(user),
    })


# ===========================================================================
# SESSIONS
# ===========================================================================

@app.route("/api/sessions")
@login_required
def assign_pics_to_session(session_id):
    """Assign PICs (divisions) to a session"""
    if current_user.role not in ['admin', 'ketua', 'pembina']:
        abort(403)
    
    session = Session.query.get_or_404(session_id)
    
    if request.method == 'POST':
        # Get selected PIC IDs
        pic_ids = request.form.getlist('pic_ids')
        
        try:
            # Remove all existing PIC assignments for this session
            SessionPIC.query.filter_by(session_id=session_id).delete()
            
            # Add new PIC assignments
            for pic_id_str in pic_ids:
                try:
                    pic_id = int(pic_id_str)
                    # Verify PIC exists
                    pic = Pic.query.get(pic_id)
                    if pic:
                        session_pic = SessionPIC(
                            session_id=session_id,
                            pic_id=pic_id
                        )
                        db.session.add(session_pic)
                except (ValueError, TypeError):
                    continue
            
            db.session.commit()
            flash(f'PICs updated for "{session.name}"', 'success')
            return redirect(url_for('assign_pics_to_session', session_id=session_id))
            
        except Exception as e:
            db.session.rollback()
            flash(f'Error updating PICs: {str(e)}', 'error')
            return redirect(url_for('assign_pics_to_session', session_id=session_id))
    
    # GET request - show form
    available_pics = Pic.query.all()
    return render_template(
        'assign_pics_to_session.html',
        session=session,
        available_pics=available_pics
    )


@app.route('/session/<int:session_id>/remove-pic/<int:pic_id>', methods=['POST'])
@login_required
def delete_session(session_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    s = Session.query.get_or_404(session_id)
    name = s.name
    try:
        SessionPIC.query.filter_by(session_id=session_id).delete()
        Attendance.query.filter_by(session_id=session_id).delete()
        Notulensi.query.filter_by(session_id=session_id).delete()
        db.session.delete(s)
        db.session.commit()
        return jsonify({"success": True, "message": f'Session "{name}" deleted'})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/sessions/<int:session_id>/lock", methods=["POST"])
@login_required
def lock_session(session_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    s = Session.query.get_or_404(session_id)
    s.is_locked = True
    db.session.commit()
    return jsonify({"success": True, "is_locked": True, "session": serialize_session(s)})


@app.route("/api/sessions/<int:session_id>/status")
@login_required
def get_session_status(session_id):
    s = Session.query.get_or_404(session_id)
    return jsonify({"success": True, "is_locked": s.is_locked, "session_id": s.id, "name": s.name})


@app.route("/api/sessions/<int:session_id>/pics", methods=["GET"])
@login_required
def get_session_pics(session_id):
    s = Session.query.get_or_404(session_id)
    return jsonify({
        "success": True,
        "session_id": session_id,
        "assigned_pics": [{"id": p.id, "name": p.name, "description": p.description} for p in s.assigned_pics],
    })


@app.route("/api/sessions/<int:session_id>/pics", methods=["PUT"])
@login_required
def assign_pics_to_session(session_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    s = Session.query.get_or_404(session_id)
    data = request.get_json() or {}
    pic_ids = data.get("pic_ids", [])

    try:
        SessionPIC.query.filter_by(session_id=session_id).delete()
        for pid in pic_ids:
            pic = Pic.query.get(pid)
            if pic:
                db.session.add(SessionPIC(session_id=session_id, pic_id=pid))
        db.session.commit()
        return jsonify({"success": True, "message": "PICs updated", "session": serialize_session(s)})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/sessions/<int:session_id>/pics/<int:pic_id>", methods=["DELETE"])
@login_required
def remove_pic_from_session(session_id, pic_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    sp = SessionPIC.query.filter_by(session_id=session_id, pic_id=pic_id).first()
    if not sp:
        return jsonify({"success": False, "message": "PIC assignment not found"}), 404

    name = sp.pic.name
    db.session.delete(sp)
    db.session.commit()
    return jsonify({"success": True, "message": f"Removed {name} from session"})


# ===========================================================================
# ATTENDANCE
# ===========================================================================

@app.route("/api/attendance", methods=["POST"])
@login_required
def api_attendance():
    data = request.get_json() or {}
    session_id = data.get("session_id")
    user_id = data.get("user_id")
    status = data.get("status")

    if not all([session_id, user_id, status]):
        return jsonify({"success": False, "error": "invalid_data", "message": "Missing required fields"}), 400

    try:
        session_id, user_id = int(session_id), int(user_id)
    except (ValueError, TypeError):
        return jsonify({"success": False, "error": "invalid_data", "message": "Invalid ID format"}), 400

    s = Session.query.get(session_id)
    if not s:
        return jsonify({"success": False, "error": "not_found", "message": "Session not found"}), 404
    if s.is_locked:
        return jsonify({"success": False, "error": "session_locked", "message": "Session is locked"}), 403
    if not can_mark_attendance(current_user, s.pic_id):
        return jsonify({"success": False, "error": "forbidden", "message": "No permission to mark attendance"}), 403

    existing = Attendance.query.filter_by(session_id=session_id, user_id=user_id, attendance_type="regular").first()
    if existing:
        return jsonify({"success": False, "error": "already_marked", "message": "Attendance already recorded"}), 409

    wib = timezone(timedelta(hours=7))
    att = Attendance(session_id=session_id, user_id=user_id, status=status, attendance_type="regular", timestamp=datetime.now(wib))
    try:
        db.session.add(att)
        db.session.commit()
        return jsonify({"success": True, "attendance": serialize_attendance(att)}), 201
    except IntegrityError:
        db.session.rollback()
        return jsonify({"success": False, "error": "already_marked", "message": "Attendance already recorded"}), 409
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": "database_error", "message": str(e)}), 500


@app.route("/api/attendance/core", methods=["POST"])
@login_required
def api_attendance_core():
    if not is_core_user(current_user):
        return jsonify({"success": False, "error": "forbidden", "message": "Access denied"}), 403

    data = request.get_json() or {}
    session_id = data.get("session_id")
    user_id = data.get("user_id")
    status = data.get("status")

    if not all([session_id, user_id, status]):
        return jsonify({"success": False, "error": "invalid_data", "message": "Missing required fields"}), 400

    try:
        session_id, user_id = int(session_id), int(user_id)
    except (ValueError, TypeError):
        return jsonify({"success": False, "error": "invalid_data", "message": "Invalid ID format"}), 400

    s = Session.query.get(session_id)
    if not s:
        return jsonify({"success": False, "error": "not_found", "message": "Session not found"}), 404
    if s.is_locked:
        return jsonify({"success": False, "error": "session_locked", "message": "Session is locked"}), 403

    target = User.query.get(user_id)
    if not target or not is_core_user(target):
        return jsonify({"success": False, "error": "not_core_user", "message": "User is not a core member"}), 400

    existing = Attendance.query.filter_by(session_id=session_id, user_id=user_id, attendance_type="core").first()
    if existing:
        return jsonify({"success": False, "error": "already_marked", "message": "Attendance already recorded"}), 409

    wib = timezone(timedelta(hours=7))
    att = Attendance(session_id=session_id, user_id=user_id, status=status, attendance_type="core", timestamp=datetime.now(wib))
    try:
        db.session.add(att)
        db.session.commit()
        return jsonify({"success": True, "attendance": serialize_attendance(att)}), 201
    except IntegrityError:
        db.session.rollback()
        return jsonify({"success": False, "error": "already_marked", "message": "Attendance already recorded"}), 409
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": "database_error", "message": str(e)}), 500


@app.route("/api/attendance/history")
@login_required
def attendance_history():
    """Current user's own attendance history."""
    records = Attendance.query.filter_by(user_id=current_user.id).all()
    summary = {
        "present": sum(1 for r in records if r.status == "present"),
        "absent": sum(1 for r in records if r.status == "absent"),
        "excused": sum(1 for r in records if r.status == "excused"),
        "late": sum(1 for r in records if r.status == "late"),
        "total": len(records),
    }
    return jsonify({"success": True, "records": [serialize_attendance(r) for r in records], "summary": summary})


@app.route("/api/attendance/history/all")
@login_required
def attendance_history_all():
    """Admin: list of members for drill-down. Returns member summaries."""
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    users = User.query.filter_by(role="member").order_by(User.name).all()
    return jsonify({"success": True, "members": [serialize_user(u) for u in users]})


@app.route("/api/attendance/history/<int:user_id>")
@login_required
def attendance_history_for_user(user_id):
    """Admin: full attendance history for a specific member."""
    if current_user.role not in ["admin", "ketua", "pembina"] and current_user.id != user_id:
        return jsonify({"success": False, "message": "Access denied"}), 403

    user = User.query.get_or_404(user_id)
    records = Attendance.query.filter_by(user_id=user_id).all()
    summary = {
        "present": sum(1 for r in records if r.status == "present"),
        "absent": sum(1 for r in records if r.status == "absent"),
        "excused": sum(1 for r in records if r.status == "excused"),
        "late": sum(1 for r in records if r.status == "late"),
        "total": len(records),
    }
    return jsonify({
        "success": True,
        "user": serialize_user(user),
        "records": [serialize_attendance(r) for r in records],
        "summary": summary,
    })


@app.route("/api/export/attendance/<int:session_id>")
@login_required
def export_attendance(session_id):
    """Returns a .docx file — keep as binary response."""
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    s = Session.query.get_or_404(session_id)
    records = (
        db.session.query(Attendance, User.name, User.email, User.role)
        .join(User, Attendance.user_id == User.id)
        .filter(Attendance.session_id == session_id)
        .order_by(User.name)
        .all()
    )

    if not records:
        return jsonify({"success": False, "message": "No attendance records found"}), 404

    wib = timezone(timedelta(hours=7))
    doc = Document()
    doc.add_heading(f"Attendance Report: {s.name}", 0)
    doc.add_paragraph(f"Date: {s.date}")
    doc.add_paragraph(f"Total Attendees: {len(records)}")
    doc.add_paragraph("")

    summary = {k: sum(1 for a, _, _, _ in records if a.status == k) for k in ("present", "absent", "excused", "late")}
    doc.add_heading("Summary", level=1)
    st = doc.add_table(rows=5, cols=2)
    st.style = "Light Grid Accent 1"
    for i, (label, val) in enumerate([("Status", "Count"), ("Present", str(summary["present"])), ("Absent", str(summary["absent"])), ("Excused", str(summary["excused"])), ("Late", str(summary["late"]))]):
        st.rows[i].cells[0].text = label
        st.rows[i].cells[1].text = val

    doc.add_paragraph("")
    doc.add_heading("Detailed Records", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Name", "Role", "Status", "Time", "Type"]):
        hdr[i].text = h

    for att, name, email, role in records:
        row = table.add_row().cells
        row[0].text = name
        row[1].text = role.capitalize()
        row[2].text = att.status.capitalize()
        row[3].text = att.timestamp.astimezone(wib).strftime("%H:%M") if att.timestamp else ""
        row[4].text = att.attendance_type.capitalize()

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    filename = f"attendance_{s.name.replace(' ', '_')}_{s.date}.docx"
    return Response(
        bio,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


# ===========================================================================
# PICs (DIVISIONS)
# ===========================================================================

@app.route("/api/pics")
@login_required
def list_pics():
    pics = Pic.query.order_by(Pic.name).all()
    return jsonify({"success": True, "pics": [serialize_pic(p) for p in pics]})


@app.route("/api/pics", methods=["POST"])
@login_required
def create_pic():
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    data = request.get_json() or {}
    name = data.get("name", "").strip()
    description = data.get("description", "").strip() or None

    if not name:
        return jsonify({"success": False, "message": "PIC name is required"}), 400
    if Pic.query.filter_by(name=name).first():
        return jsonify({"success": False, "message": f"PIC '{name}' already exists"}), 409

    pic = Pic(name=name, description=description)
    db.session.add(pic)
    db.session.commit()
    return jsonify({"success": True, "message": f"PIC '{name}' created", "pic": serialize_pic(pic)}), 201


@app.route("/api/pics/<int:pic_id>", methods=["DELETE"])
@login_required
def delete_pic(pic_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    pic = Pic.query.get_or_404(pic_id)
    for user in pic.members:
        user.pic_id = None
        user.can_mark_attendance = False
    SessionPIC.query.filter_by(pic_id=pic_id).delete()
    db.session.delete(pic)
    db.session.commit()
    return jsonify({"success": True, "message": f"PIC '{pic.name}' deleted"})


# ===========================================================================
# NOTULENSI
# ===========================================================================

@app.route("/api/notulensi")
@login_required
def list_notulensi():
    sessions = Session.query.order_by(Session.date.desc()).all()
    notes = {n.session_id: n for n in Notulensi.query.all()}
    result = []
    for s in sessions:
        note = notes.get(s.id)
        result.append({
            "session_id": s.id,
            "session_name": s.name,
            "session_date": s.date,
            "has_notulensi": note is not None,
            "notulensi": serialize_notulensi(note) if note else None,
        })
    return jsonify({"success": True, "items": result})


@app.route("/api/notulensi/<int:session_id>", methods=["GET"])
@login_required
def get_notulensi(session_id):
    s = Session.query.get_or_404(session_id)
    note = Notulensi.query.filter_by(session_id=session_id).first()
    return jsonify({
        "success": True,
        "session": serialize_session(s),
        "notulensi": serialize_notulensi(note) if note else None,
        "can_edit": current_user.role in ["admin", "ketua", "pembina"],
    })


@app.route("/api/notulensi/<int:session_id>", methods=["POST"])
@login_required
def save_notulensi(session_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    data = request.get_json() or {}
    content = data.get("content", "").strip()
    if not content or content in ["<p><br></p>", "<p></p>"]:
        return jsonify({"success": False, "message": "Content cannot be empty"}), 400

    note = Notulensi.query.filter_by(session_id=session_id).first()
    if note:
        note.content = content
        note.updated_at = datetime.utcnow()
    else:
        note = Notulensi(session_id=session_id, content=content)
        db.session.add(note)

    db.session.commit()
    return jsonify({"success": True, "notulensi": serialize_notulensi(note)})


@app.route("/api/notulensi/by-id/<int:notulensi_id>", methods=["DELETE"])
@login_required
def delete_notulensi(notulensi_id):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    note = Notulensi.query.get_or_404(notulensi_id)
    db.session.delete(note)
    db.session.commit()
    return jsonify({"success": True, "message": "Notulensi deleted"})


# ===========================================================================
# CALENDAR & NEWS FEED
# ===========================================================================

ISLAMIC_HOLIDAYS = {
    "01-01": "Islamic New Year",
    "01-09": "Day of Tasua",
    "01-10": "Day of Ashura",
    "03-12": "Mawlid al-Nabi",
    "07-01": "Start of Rajab",
    "07-27": "Isra and Mi'raj",
    "08-15": "Mid-Sha'ban (Laylat al-Bara'ah)",
    "09-01": "Start of Ramadan",
    "09-17": "Nuzul al-Qur'an",
    "09-21": "Laylat al-Qadr (possible)",
    "09-23": "Laylat al-Qadr (possible)",
    "09-25": "Laylat al-Qadr (possible)",
    "09-27": "Laylat al-Qadr (possible)",
    "09-29": "Laylat al-Qadr (possible)",
    "10-01": "Eid al-Fitr",
    "10-02": "Eid al-Fitr (Day 2)",
    "11-01": "Start of Dhu al-Qi'dah",
    "12-01": "Start of Dhu al-Hijjah",
    "12-08": "Day of Tarwiyah",
    "12-09": "Day of Arafah",
    "12-10": "Eid al-Adha",
    "12-11": "Days of Tashreeq",
    "12-12": "Days of Tashreeq",
    "12-13": "Days of Tashreeq",
}


def get_hijri_date(gregorian_date_str):
    try:
        g = datetime.strptime(gregorian_date_str, "%Y-%m-%d").date()
        h = HijriDate(g.year, g.month, g.day, gr=True)
        return f"{h.day} {h.month_name} {h.year} H"
    except Exception:
        return ""


def get_hijri_key(g_date):
    h = HijriDate(g_date.year, g_date.month, g_date.day, gr=True)
    return f"{h.month:02d}-{h.day:02d}", h


@app.route("/api/calendar")
@login_required
def calendar_events():
    events = []
    for s in Session.query.all():
        hijri = get_hijri_date(s.date)
        events.append({
            "title": f"{s.name} ({hijri})",
            "start": s.date,
            "extendedProps": {"type": "rohis_session", "session_id": s.id},
        })

    today = date.today()
    current = date(today.year - 1, 1, 1)
    end = date(today.year + 1, 12, 31)
    while current <= end:
        key, h = get_hijri_key(current)
        if key in ISLAMIC_HOLIDAYS:
            events.append({
                "title": f"{ISLAMIC_HOLIDAYS[key]} ({h.day} {h.month_name} {h.year} H)",
                "start": current.isoformat(),
                "allDay": True,
                "backgroundColor": "#1e88e5",
                "borderColor": "#1565c0",
                "textColor": "#ffffff",
                "extendedProps": {
                    "type": "islamic_holiday",
                    "hijri": f"{h.day} {h.month_name} {h.year} H",
                },
            })
        current = current.fromordinal(current.toordinal() + 1)

    return jsonify(events)


@app.route("/api/feed")
@login_required
def news_feed():
    try:
        today_str = str(date.today())
        upcoming = Session.query.filter(Session.date >= today_str).order_by(Session.date.asc()).limit(3).all()
        recent = (
            db.session.query(Notulensi, Session)
            .join(Session, Notulensi.session_id == Session.id)
            .order_by(Notulensi.updated_at.desc())
            .limit(3)
            .all()
        )

        upcoming_data = []
        for s in upcoming:
            pics = Pic.query.join(SessionPIC, Pic.id == SessionPIC.pic_id).filter(SessionPIC.session_id == s.id).all()
            upcoming_data.append({
                "id": s.id,
                "name": s.name,
                "date": s.date,
                "pic": ", ".join(p.name for p in pics) if pics else "No PIC assigned",
            })

        recent_data = []
        for note, s in recent:
            summary = "Meeting notes available."
            if note.content:
                try:
                    summary = summarize_notulensi(note.content) if os.environ.get("GROQ_API_KEY") else _plain_preview(note.content)
                except Exception:
                    summary = _plain_preview(note.content)
            recent_data.append({
                "id": note.id,
                "session_name": s.name,
                "session_date": s.date,
                "summary": summary,
                "updated_at": (note.updated_at or note.created_at).strftime("%d %b %Y"),
            })

        return jsonify({"success": True, "upcoming": upcoming_data, "recent": recent_data})
    except Exception as e:
        logger.exception("News feed error")
        return jsonify({"success": True, "upcoming": [], "recent": [], "error": str(e)})


def _plain_preview(html_content, max_len=150):
    import re
    from html import unescape
    text = unescape(re.sub("<[^<]+?>", "", html_content)).strip()
    return (text[:max_len] + "...") if len(text) > max_len else (text or "Meeting notes available.")


# ===========================================================================
# CHATBOT
# ===========================================================================

@app.route('/api/cron/piket-reminder', methods=['POST'])
@app.route('/api/cron/piket-reminder', methods=['POST'])
def cron_piket_reminder():
    expected_token = os.environ.get('CRON_SECRET_TOKEN')
    
    if not expected_token:
        app.logger.error("CRON_SECRET_TOKEN not set - piket reminders disabled")
        return jsonify({
            'success': False,
            'error': 'Service not configured'
        }), 503 
    try:
        # Get current day in WIB timezone (UTC+7)
        wib = timezone(timedelta(hours=7))
        now_wib = datetime.now(wib)
        
        # Get day of week (0=Monday, 6=Sunday)
        day_of_week = now_wib.weekday()
        day_names = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
        day_name = day_names[day_of_week]
        date_str = now_wib.strftime('%d %B %Y')
        
        # Find jadwal for today
        jadwal = JadwalPiket.query.filter_by(day_of_week=day_of_week).first()
        
        if not jadwal:
            # No jadwal configured for this day - log and return
            log = EmailReminderLog(
                day_of_week=day_of_week,
                day_name=day_name,
                recipients_count=0,
                recipients='[]',
                status='skipped',
                error_message='No jadwal piket configured for this day'
            )
            db.session.add(log)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'No piket scheduled for {day_name}',
                'day': day_name,
                'date': date_str,
                'recipients_count': 0
            }), 200
        
        # Get assigned users for today
        assignments = PiketAssignment.query.filter_by(jadwal_id=jadwal.id).all()
        
        if not assignments:
            # Jadwal exists but no members assigned
            log = EmailReminderLog(
                day_of_week=day_of_week,
                day_name=day_name,
                recipients_count=0,
                recipients='[]',
                status='skipped',
                error_message='No members assigned to this day'
            )
            db.session.add(log)
            db.session.commit()
            
            return jsonify({
                'success': True,
                'message': f'No members assigned for {day_name}',
                'day': day_name,
                'date': date_str,
                'recipients_count': 0
            }), 200
        
        # Collect recipient emails
        recipients = []
        for assignment in assignments:
            user = assignment.user
            if user and user.email:
                recipients.append(user.email)
        
        if not recipients:
            # Assignments exist but no valid emails
            log = EmailReminderLog(
                day_of_week=day_of_week,
                day_name=day_name,
                recipients_count=0,
                recipients='[]',
                status='failed',
                error_message='No valid email addresses found for assigned members'
            )
            db.session.add(log)
            db.session.commit()
            
            return jsonify({
                'success': False,
                'error': 'No valid email addresses found',
                'day': day_name,
                'date': date_str
            }), 500
        
        # Send emails
        email_service = get_email_service()
        result = email_service.send_piket_reminder(
            recipients=recipients,
            day_name=day_name,
            date_str=date_str,
            additional_info=""  # Can be customized
        )
        
        # Log the reminder
        log = EmailReminderLog(
            day_of_week=day_of_week,
            day_name=day_name,
            recipients_count=len(recipients),
            recipients=json.dumps(recipients),
            status='success' if result['success'] else 'partial',
            error_message=result.get('message') if not result['success'] else None
        )
        db.session.add(log)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'message': result['message'],
            'day': day_name,
            'date': date_str,
            'recipients_count': len(recipients),
            'failed_emails': result.get('failed_emails', [])
        }), 200
        
    except Exception as e:
        # Log the error
        try:
            error_log = EmailReminderLog(
                day_of_week=now_wib.weekday() if 'now_wib' in locals() else -1,
                day_name='Unknown',
                recipients_count=0,
                recipients='[]',
                status='failed',
                error_message=str(e)
            )
            db.session.add(error_log)
            db.session.commit()
        except:
            pass
        
        print(f"Piket reminder error: {e}")
        import traceback
        traceback.print_exc()
        
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ============================================================================
# ADMIN ROUTES - Manage jadwal piket
# ============================================================================

@app.route('/admin/jadwal-piket', methods=['GET', 'POST'])
@login_required
def update_piket():
    """Admin: update piket assignments for a day."""
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    data = request.get_json() or {}
    day_of_week = data.get("day_of_week")
    user_ids = data.get("user_ids", [])

    if day_of_week is None or not (0 <= int(day_of_week) <= 6):
        return jsonify({"success": False, "message": "Invalid day_of_week (0–6)"}), 400

    day_of_week = int(day_of_week)
    day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]

    try:
        jadwal = JadwalPiket.query.filter_by(day_of_week=day_of_week).first()
        if not jadwal:
            jadwal = JadwalPiket(day_of_week=day_of_week, day_name=day_names[day_of_week])
            db.session.add(jadwal)
            db.session.flush()

        PiketAssignment.query.filter_by(jadwal_id=jadwal.id).delete()
        for uid in user_ids:
            if uid:
                db.session.add(PiketAssignment(jadwal_id=jadwal.id, user_id=int(uid)))

        jadwal.updated_at = datetime.utcnow()
        db.session.commit()
        return jsonify({"success": True, "message": f"Piket for {day_names[day_of_week]} updated"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "message": str(e)}), 500


@app.route("/api/piket/<int:day_of_week>", methods=["DELETE"])
@login_required
def clear_piket(day_of_week):
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    jadwal = JadwalPiket.query.filter_by(day_of_week=day_of_week).first()
    if not jadwal:
        return jsonify({"success": False, "message": "No schedule found for that day"}), 404

    PiketAssignment.query.filter_by(jadwal_id=jadwal.id).delete()
    db.session.commit()
    return jsonify({"success": True, "message": "Assignments cleared"})


@app.route("/api/piket/logs")
@login_required
def piket_logs():
    if current_user.role not in ["admin", "ketua", "pembina"]:
        return jsonify({"success": False, "message": "Access denied"}), 403

    logs = EmailReminderLog.query.order_by(EmailReminderLog.sent_at.desc()).limit(100).all()
    result = []
    for log in logs:
        result.append({
            "id": log.id,
            "day_of_week": log.day_of_week,
            "day_name": log.day_name,
            "recipients_count": log.recipients_count,
            "recipients": json.loads(log.recipients) if log.recipients else [],
            "sent_at": log.sent_at.isoformat() if log.sent_at else None,
            "status": log.status,
            "error_message": log.error_message,
        })
    return jsonify({"success": True, "logs": result})


@app.route("/api/piket/test", methods=["POST"])
@login_required
def test_piket_reminder():
    if current_user.role not in ["admin"]:
        return jsonify({"success": False, "message": "Admin only"}), 403

    data = request.get_json() or {}
    wib = timezone(timedelta(hours=7))
    day_of_week = int(data.get("day_of_week", datetime.now(wib).weekday()))
    day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
    day_name = day_names[day_of_week]

    jadwal = JadwalPiket.query.filter_by(day_of_week=day_of_week).first()
    if not jadwal or not jadwal.assignments:
        return jsonify({"success": False, "message": f"No assignments for {day_name}"}), 404

    recipients = [a.user.email for a in jadwal.assignments if a.user and a.user.email]
    if not recipients:
        return jsonify({"success": False, "message": "No valid email addresses found"}), 404

    result = get_email_service().send_piket_reminder(
        recipients=recipients,
        day_name=day_name,
        date_str=datetime.now().strftime("%d %B %Y"),
        additional_info="⚠️ This is a TEST reminder from the admin panel.",
    )
    return jsonify({"success": result["success"], "message": result["message"], "failed_emails": result.get("failed_emails", [])})


# ===========================================================================
# CRON (called by external scheduler e.g. cron-job.org)
# ===========================================================================

@app.route("/api/cron/piket-reminder", methods=["POST"])
def cron_piket_reminder():
    expected = os.environ.get("CRON_SECRET_TOKEN")
    if not expected:
        return jsonify({"success": False, "error": "Service not configured"}), 503

    provided = request.headers.get("X-Cron-Secret") or (request.get_json() or {}).get("secret")
    if not provided or provided != expected:
        return jsonify({"success": False, "error": "Unauthorized"}), 401

    try:
        wib = timezone(timedelta(hours=7))
        now_wib = datetime.now(wib)
        day_of_week = now_wib.weekday()
        day_names = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
        day_name = day_names[day_of_week]
        date_str = now_wib.strftime("%d %B %Y")

        def _log(status, error=None, count=0, recipients="[]"):
            log = EmailReminderLog(
                day_of_week=day_of_week, day_name=day_name,
                recipients_count=count, recipients=recipients,
                status=status, error_message=error,
            )
            db.session.add(log)
            db.session.commit()

        jadwal = JadwalPiket.query.filter_by(day_of_week=day_of_week).first()
        if not jadwal:
            _log("skipped", "No jadwal piket configured for this day")
            return jsonify({"success": True, "message": f"No piket for {day_name}", "recipients_count": 0})

        assignments = PiketAssignment.query.filter_by(jadwal_id=jadwal.id).all()
        if not assignments:
            _log("skipped", "No members assigned")
            return jsonify({"success": True, "message": f"No members for {day_name}", "recipients_count": 0})

        recipients = [a.user.email for a in assignments if a.user and a.user.email]
        if not recipients:
            _log("failed", "No valid emails")
            return jsonify({"success": False, "error": "No valid emails"}), 500

        result = get_email_service().send_piket_reminder(recipients=recipients, day_name=day_name, date_str=date_str)
        _log(
            "success" if result["success"] else "partial",
            result.get("message") if not result["success"] else None,
            len(recipients),
            json.dumps(recipients),
        )
        return jsonify({
            "success": True,
            "message": result["message"],
            "day": day_name,
            "date": date_str,
            "recipients_count": len(recipients),
            "failed_emails": result.get("failed_emails", []),
        })

    except Exception as e:
        logger.exception("Cron reminder error")
        return jsonify({"success": False, "error": str(e)}), 500


# ===========================================================================
# ERROR HANDLERS
# ===========================================================================

@app.errorhandler(403)
def forbidden(_):
    return jsonify({"success": False, "error": "forbidden", "message": "You do not have permission to perform this action"}), 403


@app.errorhandler(404)
def not_found(_):
    return jsonify({"success": False, "error": "not_found", "message": "The requested resource was not found"}), 404


@app.errorhandler(405)
def method_not_allowed(_):
    return jsonify({"success": False, "error": "method_not_allowed"}), 405


@app.errorhandler(500)
def internal_error(e):
    return jsonify({"success": False, "error": "internal_server_error", "message": str(e)}), 500


# ===========================================================================
# Entry point
# ===========================================================================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)